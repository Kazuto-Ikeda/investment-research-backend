from fastapi import Body, Query, BackgroundTasks, HTTPException
from fastapi.responses import FileResponse
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, List, Optional
import os
import re
import logging
import mistune
from mistune import Markdown
from mistune import create_markdown
from mistune.plugins import plugin_table
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ############### 追加部分：カスタムバレット番号定義の追加 ###############
def add_custom_bullet_numbering(document, num_id=99):
    """
    カスタムの箇条書き番号定義を追加します。
    各インデントレベルごとに異なるバレット記号を設定します。
    ここではレベル 0: "•", レベル 1: "◦", レベル 2: "▪" としています。
    """
    # ### 変更箇所：内部XML要素 _element を利用して取得
    numbering = document.part.numbering_part._element

    # abstractNum 要素の作成
    abstract_num = OxmlElement('w:abstractNum')
    abstract_num.set(qn('w:abstractNumId'), '0')  # 固定の abstractNumId "0"

    # 各レベルの設定（レベル 0～2 の例）
    bullet_symbols = ['•', '◦', '▪']  # 追加するバレット記号
    for lvl in range(3):
        lvl_element = OxmlElement('w:lvl')
        lvl_element.set(qn('w:ilvl'), str(lvl))

        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl_element.append(start)

        num_fmt = OxmlElement('w:numFmt')
        num_fmt.set(qn('w:val'), 'bullet')  # bullet 表示にする
        lvl_element.append(num_fmt)

        lvl_text = OxmlElement('w:lvlText')
        lvl_text.set(qn('w:val'), bullet_symbols[lvl])
        lvl_element.append(lvl_text)

        lvl_jc = OxmlElement('w:lvlJc')
        lvl_jc.set(qn('w:val'), 'left')
        lvl_element.append(lvl_jc)

        # インデント設定（例：レベル毎にインデントを増加）
        pPr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), str(720 + 360 * lvl))  # 720 + 360*lvl
        ind.set(qn('w:hanging'), '360')
        pPr.append(ind)
        lvl_element.append(pPr)

        abstract_num.append(lvl_element)

    # abstractNum を numbering に追加（直接append()）
    numbering.append(abstract_num)
    abstract_num_id = abstract_num.get(qn('w:abstractNumId'))

    # num 要素の作成
    num = OxmlElement('w:num')
    num.set(qn('w:numId'), str(num_id))
    abstract_num_ref = OxmlElement('w:abstractNumId')
    abstract_num_ref.set(qn('w:val'), str(abstract_num_id))
    num.append(abstract_num_ref)
    numbering.append(num)
# ############### 追加部分ここまで ###############


class DocxRenderer(mistune.AstRenderer):
    """
    Mistune 2.x の AstRenderer を継承し、Markdown の AST を Python‑docx で Word に変換するクラス。
    """
    def __init__(self, document: Document):
        super().__init__()
        self.document = document
        self.current_paragraph = None  # 処理中の段落オブジェクト
        self._reset_numbering_for_next_list = False

    def render(self, tokens, state):
        """
        AST全体を走査し、各ノードに応じた描画メソッドを呼び出す。
        """
        for token in tokens:
            node_type = token['type']
            if node_type == 'heading':
                self._render_heading(token)
            elif node_type == 'paragraph':
                self.current_paragraph = self.document.add_paragraph()
                inline_children = token.get('children', [])
                self._render_inline_children(inline_children)
            elif node_type == 'list':
                self._render_list(token, level=1)
            elif node_type == 'table':
                self._render_table(token)
            else:
                logging.debug(f"[DocxRenderer] Skip unknown node: {node_type}")
        return ''

    #######################################################
    # heading (見出し)
    #######################################################
    def _render_heading(self, token):
        """
        見出しを描画し、見出し直後のorderedリストで番号リセットするためのフラグを立てる。
        """
        level = token['level']
        children = token.get('children', [])
        self.current_paragraph = self.document.add_paragraph()

        if level == 1:
            fsize = Pt(18)
            self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif level == 2:
            fsize = Pt(16)
            self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            fsize = Pt(14)
            self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        for child in children:
            txt = self._extract_text(child)
            run = self.current_paragraph.add_run(txt)
            run.font.size = fsize
            run.bold = True

        # 見出し直後のorderedリストで番号をリセットするためのフラグ
        self._reset_numbering_for_next_list = True

    #######################################################
    # paragraph (段落)
    #######################################################
    def _render_paragraph(self, token):
        """
        段落を描画する。
        """
        logging.debug(f"_render_paragraph: {token}")
        self.current_paragraph = self.document.add_paragraph()
        inline_children = token.get('children', [])
        self._render_inline_children(inline_children)

    def _render_inline_children(self, children):
        for node in children:
            ntype = node['type']
            if ntype == 'text':
                run = self.current_paragraph.add_run(node['text'])
            elif ntype == 'strong':
                txt = self._extract_text(node)
                run = self.current_paragraph.add_run(txt)
                run.bold = True
            elif ntype == 'emphasis':
                txt = self._extract_text(node)
                run = self.current_paragraph.add_run(txt)
                run.italic = True
            else:
                txt = self._extract_text(node)
                run = self.current_paragraph.add_run(txt)

    #######################################################
    # list + list_item (入れ子対応)
    #######################################################
    def _render_list(self, token, level=1):
        """
        リスト全体を描画する。ordered・unorderedに応じたスタイル設定を行う。
        """
        ordered = token.get('ordered', False)
        for child in token.get('children', []):
            if child['type'] == 'list_item':
                self._render_list_item(child, ordered, level)
            elif child['type'] == 'list':
                self._render_list(child, level=level+1)

    # ############### 追加部分： カスタムバレット適用用の静的メソッド ###############
    @staticmethod
    def _apply_custom_bullet(paragraph, level=0, num_id=99):
        """
        指定した段落に対して、カスタム箇条書き番号定義（num_id=99）を適用し、
        インデントレベルに応じたバレット記号を設定します。
        """
        p = paragraph._p  # 内部のXMLオブジェクト
        pPr = p.get_or_add_pPr()
        numPr = pPr.get_or_add_numPr()
        for child in list(numPr):
            numPr.remove(child)
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), str(level))
        numPr.append(ilvl)
        numId_elm = OxmlElement('w:numId')
        numId_elm.set(qn('w:val'), str(num_id))
        numPr.append(numId_elm)
    # ############### 追加部分ここまで ###############

    def _render_list_item(self, token, ordered, level):
        """
        list_item を描画する。
        ・orderedの場合:
           - 見出し直後の最初のorderedリストでは _restart_numbering() を利用して番号を再スタートさせるが、
             このとき、レベルは現在のリストのインデントレベル (level) を指定する（※変更点）。
           - その後は、_apply_custom_bullet() を用いて、現在のレベルに応じたバレット記号を適用する。
        ・unorderedの場合は List Bullet スタイルが適用される。
        """
        if ordered:
            style = 'List Number'
        else:
            style = 'List Bullet'
        self.current_paragraph = self.document.add_paragraph(style=style)
        self.current_paragraph.paragraph_format.left_indent = Cm(0.5 * level)

        if ordered:
            if self._reset_numbering_for_next_list:
                # ★★【変更点】★★: レベルを固定0ではなく、現在の level を指定する
                self._restart_numbering(self.current_paragraph, level=level, num_id=1)
                self._reset_numbering_for_next_list = False
            else:
                DocxRenderer._apply_custom_bullet(self.current_paragraph, level=level, num_id=99)
        # unorderedの場合はそのまま

        for child in token.get('children', []):
            ctype = child['type']
            if ctype == 'paragraph':
                inline_nodes = child.get('children', [])
                self._render_inline_children(inline_nodes)
            elif ctype == 'list':
                self._render_list(child, level=level+1)
            else:
                txt = self._extract_text(child)
                self.current_paragraph.add_run(txt)

    # ############### 追加部分： _restart_numbering の復活 ###############
    def _restart_numbering(self, paragraph, level=0, num_id=1):
        """
        指定した段落に対して、番号付けを num_id の定義に沿って設定し、
        レベル (ilvl) も設定する。これにより、その段落の番号付けを再スタートする。
        """
        p = paragraph._p  # 段落の内部XMLオブジェクト
        pPr = p.get_or_add_pPr()
        numPr = pPr.get_or_add_numPr()
        for child in list(numPr):
            numPr.remove(child)
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), str(level))
        numPr.append(ilvl)
        numId_elm = OxmlElement('w:numId')
        numId_elm.set(qn('w:val'), str(num_id))
        numPr.append(numId_elm)
    # ############### 追加部分ここまで ###############

    #######################################################
    # blockquote (引用)
    #######################################################
    def _render_blockquote(self, token):
        """
        引用ブロックを描画する。スタイル 'Intense Quote' を適用する。
        """
        self.current_paragraph = self.document.add_paragraph(style='Intense Quote')
        for child in token.get('children', []):
            if child['type'] == 'paragraph':
                text = ''.join(self._extract_text(n) for n in child.get('children', []))
                self.current_paragraph.add_run(text)

    #######################################################
    # thematic_break (水平線)
    #######################################################
    def _render_thematic_break(self, token):
        """
        区切り線として "--------" を追加する。
        """
        hr_para = self.document.add_paragraph()
        hr_para.add_run('--------------').bold = True

    #######################################################
    # table (表)
    #######################################################
    def _render_table(self, token):
        """
        Mistune 2.x plugin_table で生成されたトークンに合わせたテーブルを描画する。
        """
        table_head = None
        table_cell = None
        for child in token.get('children', []):
            if child['type'] == 'table_head':
                table_head = child
            elif child['type'] == 'table_body':
                table_cell = child

        if table_head and 'children' in table_head:
            head_cells = table_head['children']
            col_count = len(head_cells)
        else:
            if table_cell and 'children' in table_cell and len(table_cell['children']) > 0:
                first_row = table_cell['children'][0]
                if first_row.get('type') == 'table_row':
                    col_count = len(first_row.get('children', []))
                else:
                    col_count = 0
            else:
                col_count = 0

        if col_count == 0:
            return

        table = self.document.add_table(rows=1, cols=col_count)
        table.style = 'Table Grid'

        if table_head and table_head.get('children'):
            hdr_cells = table.rows[0].cells
            for i, cell_ast in enumerate(table_head['children']):
                cell_text = self._extract_text(cell_ast)
                hdr_cells[i].text = cell_text
                for p in hdr_cells[i].paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.size = Pt(10)
        else:
            pass

        if table_cell and table_cell.get('children'):
            for row_ast in table_cell['children']:
                if row_ast.get('type') != 'table_row':
                    continue
                row_cells = table.add_row().cells
                for col_idx, cell_ast in enumerate(row_ast.get('children', [])):
                    cell_text = self._extract_text(cell_ast)
                    row_cells[col_idx].text = cell_text
        else:
            pass

    #######################################################
    # テキスト抽出用のヘルパー
    #######################################################
    def _extract_text(self, node):
        if 'text' in node:
            return node['text']
        elif 'children' in node:
            return ''.join(self._extract_text(child) for child in node['children'])
        return ''

def delete_file(path: str):
    """
    FileResponse返却後にバックグラウンドで削除するための後処理
    """
    try:
        os.remove(path)
        logging.info(f"[delete_file] Deleted file: {path}")
    except Exception as e:
        logging.error(f"[delete_file] Error: {e}")

def generate_word_file(
    background_tasks: BackgroundTasks,
    summaries: dict,            # 例: { "Perplexity": {...}, "ChatGPT": {...} }
    valuation_data: Optional[dict],
    company_name: str,
    file_name: Optional[str] = "result",
) -> FileResponse:
    """
    1. MarkdownをDocxRenderer (Mistune) でWordに変換
    2. valuation_dataを表形式で出力
    3. Wordファイルを保存し、FileResponseで返却（後処理でファイル削除）
    """
    ########################################
    # 1) Wordドキュメント作成
    ########################################
    file_name = f"{company_name}_summary_report.docx"
    document = Document()

    # ★★【変更点④】★★
    # ドキュメントのデフォルトフォントを設定（Noto Sans JP → Meiryo に変更）
    # Windowsユーザー向けに「Meiryo（メイリオ）」を利用します。
    style = document.styles['Normal']
    font = style.font
    font.name = 'Meiryo'
    font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Meiryo')
 
    # タイトル段落
    title_para = document.add_paragraph(f"{company_name} - 要約レポート")
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_para.runs[0]
    run_title.font.size = Pt(18)
    run_title.bold = True

    ########################################
    # 2) カスタムDocxRendererでMarkdownを解析
    ########################################
    # カスタム番号定義の追加（numId=99）
    add_custom_bullet_numbering(document, num_id=99)

    docx_renderer = DocxRenderer(document)
    md_parser = create_markdown(renderer=docx_renderer, plugins=[plugin_table])

    # カテゴリ見出しのマッピング（任意）
    category_mapping = {
        "Perplexity": "Perplexity 分析",
        "ChatGPT": "ChatGPT+SPEEDA 分析",
    }

    # セクション見出しのマッピング（任意）
    reverse_key_mapping = {
        "current_situation": "現状",
        "future_outlook": "将来性と課題",
        "investment_advantages": "競合と差別化",
        "investment_disadvantages": "Exit先検討",
        "value_up": "バリューアップ施策",
        "use_case": "M&A事例",
        "swot_analysis": "SWOT分析",
    }

    for main_category, sections in summaries.items():
        # カテゴリ見出し
        cat_jp = category_mapping.get(main_category, main_category)
        cat_para = document.add_paragraph(cat_jp)
        cat_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cat_run = cat_para.runs[0]
        cat_run.font.size = Pt(16)
        cat_run.bold = True

        # 各セクション
        for idx, (sec_key, sec_text) in enumerate(sections.items(), start=1):
            sec_jp = reverse_key_mapping.get(sec_key, sec_key)
            sec_para = document.add_paragraph(f"{idx}. {sec_jp}")
            sec_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            sec_run = sec_para.runs[0]
            sec_run.font.size = Pt(14)
            sec_run.bold = True

            # Markdown → Word
            tokens = md_parser.parse(sec_text or "内容がありません")
            docx_renderer.render(tokens, state={})

    ########################################
    # 3) バリュエーションテーブル追加 (オプション)
    ########################################
    if valuation_data:
        val_para = document.add_paragraph("バリュエーション")
        val_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        val_run = val_para.runs[0]
        val_run.font.size = Pt(16)
        val_run.bold = True

        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"

        hdr = table.rows[0].cells
        hdr[0].text = "項目"
        hdr[1].text = "直近実績"
        hdr[2].text = "進行期見込"
        for c_ in hdr:
            for p_ in c_.paragraphs:
                for r_ in p_.runs:
                    r_.font.bold = True
                    r_.font.size = Pt(10)

        for label_, val_obj in valuation_data.items():
            row_cells = table.add_row().cells
            row_cells[0].text = label_
            if isinstance(val_obj, dict):
                row_cells[1].text = val_obj.get("current", "不明")
                row_cells[2].text = val_obj.get("forecast", "不明")
            else:
                row_cells[1].text = str(val_obj)
                row_cells[2].text = "不明"

            for cell_ in row_cells:
                for p_ in cell_.paragraphs:
                    p_.style = document.styles['Normal']
                    for r_ in p_.runs:
                        r_.font.size = Pt(10)

    ########################################
    # 4) Wordファイル保存 & FileResponse
    ########################################
    output_dir = "output"
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, file_name)
    document.save(output_path)

    background_tasks.add_task(delete_file, output_path)

    return FileResponse(
        output_path,
        filename=file_name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )